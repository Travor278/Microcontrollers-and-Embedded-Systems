from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parents[2]
EX4_PAPER = ROOT / "EX4" / "Paper"
TEMPLATE = ROOT / "EX1" / "Paper" / "exp1_report.docx"
OUTPUT = EX4_PAPER / "exp4_report.docx"


def set_run_font(run, east_asia: str, latin: str | None = None, size: int = 11, bold: bool = False):
    run.font.name = latin or east_asia
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.bold = bold


def add_body_paragraph(cell, text: str, first_line_indent: float = 0.74, space_after: int = 0):
    p = cell.add_paragraph()
    p.style = "Normal"
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.first_line_indent = Cm(first_line_indent)
    pf.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, "仿宋_GB2312", size=11)
    return p


def add_heading(cell, text: str, size: int = 11, top_space: int = 6):
    p = cell.add_paragraph()
    p.style = "Normal"
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.first_line_indent = Cm(0)
    pf.space_before = Pt(top_space)
    pf.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, "黑体", size=size, bold=True)
    return p


def add_center_note(cell, text: str):
    p = cell.add_paragraph()
    p.style = "Normal"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.line_spacing = 1.2
    run = p.add_run(text)
    set_run_font(run, "仿宋_GB2312", size=10)
    return p


def add_simple_table(cell, headers: list[str], rows: list[list[str]]):
    table = cell.add_table(rows=1, cols=len(headers))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    for idx, text in enumerate(headers):
        table.rows[0].cells[idx].text = text
    for row_data in rows:
        row = table.add_row()
        for idx, text in enumerate(row_data):
            row.cells[idx].text = text
    for row in table.rows:
        for c in row.cells:
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    set_run_font(run, "仿宋_GB2312", size=9)
    return table


def build_report():
    EX4_PAPER.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(TEMPLATE, OUTPUT)
    doc = Document(str(OUTPUT))

    info_table = doc.tables[0]
    info_table.cell(0, 0).text = "专业 ________    班级 ________    姓名 ________    学号 ________________"
    info_table.cell(1, 0).text = "课程名称 微控制器与嵌入式系统实验              实验名称 实验四"
    info_table.cell(2, 0).text = "地点 TDX-PITE实验平台 / Proteus仿真    台号 ______    指导教师 ________    日期 ________"

    content_cell = doc.tables[1].cell(0, 0)
    content_cell.text = ""

    add_heading(content_cell, "一、实验内容、目的与要求", size=14, top_space=0)
    add_body_paragraph(content_cell, "本实验按照《2026微控制器与嵌入式系统实验任务书（适用于自动化、测控2024级）》中“实验四 静态存储器/点阵 LED 显示设计实验”的要求完成。基本题包括《单片机实验指导.pdf》“4.1 静态存储器扩展实验”、“2.5 数据排序实验”和“4.6 点阵 LED 显示设计实验”。任务书中点阵章节编号写为 4.7（4.13），但本指导书实际对应章节为 4.6。")
    add_body_paragraph(content_cell, "实验目的，一是掌握 MCS-51 单片机片内 RAM 与外部 RAM 的访问方式，理解 MOVX/XBYTE 访问外部数据存储器的特点；二是掌握数据排序程序的基本算法，能对片内 RAM 和外部 RAM 中的目标数据区进行升序排列；三是理解点阵 LED 的行列扫描原理，掌握通过字模表、行选、列数据和刷新延时实现字符显示及移动显示的方法。")
    add_body_paragraph(content_cell, "提高题要求对片内 RAM 的 40H-4FH 单元和外部 RAM 的 0000H-000FH 空间中的数据升序排列；同时，新实验箱 16x16 点阵循环显示自己的名字并通过开关控制移动方向，旧实验箱 8x8 点阵循环显示自己的学号并通过开关控制移动方向。")

    add_heading(content_cell, "二、实验硬件与软件环境条件（标注实验设备名称及设备号）")
    add_body_paragraph(content_cell, "硬件环境：PC 机、TDX-PITE 教学实验系统、TD-51 系统平台、外部静态 RAM 单元、16x16 或 8x8 点阵 LED 显示单元、方向控制开关、Proteus 虚拟仿真环境。")
    add_body_paragraph(content_cell, "软件环境：Keil μVision5、PK51 C51 工具链、Proteus 8 Professional、《2026微控制器与嵌入式系统实验任务书（适用于自动化、测控2024级）》和《单片机实验指导.pdf》。")
    add_body_paragraph(content_cell, "本实验整理的源码包括 EX4_SRAM.c、EX4_Sort_Internal.c、EX4_Sort_Internal_External_Advanced.c、EX4_DotMatrix_8x8.c 和 EX4_DotMatrix_16x16_Template.c。上述文件均已使用 C51 编译器进行语法检查，其中静态存储器、片内排序、片内/外排序、8x8 点阵和 16x16 点阵模板均可通过编译。")

    add_heading(content_cell, "三、实验线路示图、程序算法流程图")
    add_body_paragraph(content_cell, "1. 实验线路说明")
    add_body_paragraph(content_cell, "静态存储器实验中，外部 SRAM 通过数据总线、地址总线以及 RD、WR 等控制信号与单片机连接，程序通过 XBYTE[addr] 访问外部数据存储器；片内 RAM 则通过 DBYTE[addr] 访问。调试时可在 Keil Memory 窗口中同时观察 D:0x30、D:0x40 和 X:0x0000。")
    add_body_paragraph(content_cell, "点阵显示实验中，8x8 简化仿真采用 P1 作为行选信号、P2 作为列数据，P3.2 作为方向开关；16x16 新实验箱模板采用行译码和列移位输出思想，实际使用时需要根据实验箱 SM16206、SM5166 及接线图调整引脚映射，并将 name_font[] 替换为本人姓名字模。")
    add_body_paragraph(content_cell, "2. 程序算法说明")
    add_body_paragraph(content_cell, "（1）静态存储器扩展实验：先向片内 RAM 30H-3FH 写入 00H-0FH，再依次复制到外部 RAM 0000H-000FH，最后把外部 RAM 0000H-000FH 的内容读回到片内 RAM 40H-4FH。该过程验证片内 RAM 与外部 RAM 之间的数据传送。")
    add_body_paragraph(content_cell, "（2）片内 RAM 排序基本题：在 30H-39H 中写入 10 个无序数据，采用双重循环进行升序排序。外层循环确定当前位置，内层循环寻找后续更小的数据，若前一单元数据大于后一单元数据则交换，最终 30H-39H 由小到大排列。")
    add_body_paragraph(content_cell, "（3）片内/外 RAM 排序提高题：分别初始化片内 RAM 40H-4FH 和外部 RAM 0000H-000FH 两段 16 字节数据，再分别调用片内排序和外部排序过程。片内排序使用 DBYTE[] 访问目标区，外部排序使用 XBYTE[] 访问目标区。")
    add_body_paragraph(content_cell, "（4）点阵 LED 显示：程序从字模表取出当前字符的行数据，逐行选通点阵行，并向列线输出对应列数据。每行保持短暂延时后切换到下一行，利用人眼视觉暂留形成完整字符；方向开关改变列数据循环移位方向，从而实现左移或右移显示。")
    add_body_paragraph(content_cell, "实验四对应的 Mermaid 流程图已整理在《实验四算法流程图_mermaid.md》中，包括静态存储器传送、片内 RAM 排序、片内/外 RAM 排序和点阵扫描显示流程。")

    add_heading(content_cell, "四、实验调试步骤、实验数据记录及实验结果")
    add_body_paragraph(content_cell, "1. 实验调试步骤")
    add_body_paragraph(content_cell, "（1）建立 EX4 工程，分别添加 EX4_SRAM.c、EX4_Sort_Internal.c、EX4_Sort_Internal_External_Advanced.c、EX4_DotMatrix_8x8.c 或 EX4_DotMatrix_16x16_Template.c，单个工程中只保留一个 main() 源文件参与编译。")
    add_body_paragraph(content_cell, "（2）静态存储器实验进入 Debug 后，打开 Memory 窗口，分别输入 D:0x30、X:0x0000、D:0x40，运行程序后比较三段数据是否一致。")
    add_body_paragraph(content_cell, "（3）片内排序实验运行 EX4_Sort_Internal.c，在 D:0x30 处观察排序前后的数据。程序初始化数据后执行排序，停止在 while(1) 后查看 30H-39H。")
    add_body_paragraph(content_cell, "（4）提高题排序实验运行 EX4_Sort_Internal_External_Advanced.c，在 D:0x40 和 X:0x0000 两个 Memory 窗口中分别观察片内和外部 RAM 的排序结果。")
    add_body_paragraph(content_cell, "（5）点阵实验可先用 EX4_DotMatrix_8x8.c 在 Proteus 中进行简化验证：P1 接 8 行，P2 接 8 列，P3.2 接方向开关，观察字模是否循环显示、方向开关是否能改变移动方向。若使用实验箱 16x16 点阵，则根据实验箱接线修改 EX4_DotMatrix_16x16_Template.c 的引脚映射，并替换为本人姓名字模。")
    add_body_paragraph(content_cell, "2. 实验数据记录及实验结果")
    add_body_paragraph(content_cell, "（1）静态存储器实验结果如下表所示：")
    add_simple_table(
        content_cell,
        ["观察区域", "期望内容", "验证意义"],
        [
            ["D:30H-3FH", "00 01 02 ... 0F", "片内 RAM 初始化正确"],
            ["X:0000H-000FH", "00 01 02 ... 0F", "片内到外部 RAM 传送正确"],
            ["D:40H-4FH", "00 01 02 ... 0F", "外部 RAM 读回片内正确"],
        ],
    )
    add_body_paragraph(content_cell, "（2）片内 RAM 排序基本题中，初始数据为 09、11、05、31、20、16、01、1A、3F、08。排序后 D:30H-39H 应为 01、05、08、09、11、16、1A、20、31、3F，说明排序程序能正确完成升序排列。")
    add_body_paragraph(content_cell, "（3）提高题中，片内 RAM 40H-4FH 和外部 RAM 0000H-000FH 分别完成 16 字节升序排序。该结果说明同一排序思想可以作用于不同存储空间，但访问方式必须区分片内 DBYTE 和外部 XBYTE。")
    add_body_paragraph(content_cell, "（4）点阵显示中，8x8 模板能够逐行扫描显示数字字模，并根据 P3.2 方向开关改变列数据循环移位方向；16x16 模板则给出了行选、列移位、锁存和消隐的基本结构，可按本人姓名字模和实际实验箱端口进行替换。")
    add_body_paragraph(content_cell, "（5）C51 编译检查中，EX4_SRAM.c、EX4_Sort_Internal.c、EX4_Sort_Internal_External_Advanced.c、EX4_DotMatrix_8x8.c 和 EX4_DotMatrix_16x16_Template.c 均可通过语法编译，说明源码结构、头文件和 C51 存储区访问语法基本正确。")

    add_heading(content_cell, "五、实验总结与心得")
    add_body_paragraph(content_cell, "通过实验四，我进一步理解了 51 单片机对不同存储空间的访问方式。片内 RAM 访问速度快、地址空间小，外部 RAM 容量较大但需要通过总线和 MOVX 类访问完成。静态存储器实验把“写片内、传外部、再读回”的流程完整串起来，使片内/外存储空间的区别更加直观。")
    add_body_paragraph(content_cell, "排序实验说明，算法本身并不复杂，关键在于明确数据区起始地址、长度和比较交换条件。对于提高题，片内 RAM 与外部 RAM 的排序逻辑相同，但访问宏不同，因此程序中分别实现内部排序和外部排序，避免混用地址空间。")
    add_body_paragraph(content_cell, "点阵显示实验则强调扫描刷新思想。点阵的每个 LED 由行列共同决定，程序并不是一次点亮整个字符，而是快速逐行刷新。只要扫描速度足够快、字模数据正确、行列极性匹配，就能形成稳定的字符显示；改变字模偏移方向即可实现开关控制移动方向。")

    doc.save(str(OUTPUT))


if __name__ == "__main__":
    build_report()
