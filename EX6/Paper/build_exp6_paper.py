from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parents[2]
EX6_PAPER = ROOT / "EX6" / "Paper"
TEMPLATE = ROOT / "EX1" / "Paper" / "exp1_report.docx"
OUTPUT = EX6_PAPER / "exp6_report.docx"


def set_run_font(run, east_asia: str, latin: str | None = None, size: int = 11, bold: bool = False):
    run.font.name = latin or east_asia
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.bold = bold


def add_paragraph(cell, text: str, indent: float = 0.74, size: int = 11):
    p = cell.add_paragraph()
    p.style = "Normal"
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(indent)
    run = p.add_run(text)
    set_run_font(run, "仿宋_GB2312", size=size)
    return p


def add_heading(cell, text: str, level: int = 1):
    p = cell.add_paragraph()
    p.style = "Normal"
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(8 if level == 1 else 4)
    run = p.add_run(text)
    set_run_font(run, "黑体", size=14 if level == 1 else 11, bold=True)
    return p


def add_table(cell, headers: list[str], rows: list[list[str]]):
    table = cell.add_table(rows=1, cols=len(headers))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    for i, text in enumerate(headers):
        table.rows[0].cells[i].text = text
    for row_data in rows:
        row = table.add_row()
        for i, text in enumerate(row_data):
            row.cells[i].text = text
    for row in table.rows:
        for c in row.cells:
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.line_spacing = 1.15
                for run in p.runs:
                    set_run_font(run, "仿宋_GB2312", size=9)
    return table


def build_report():
    EX6_PAPER.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(TEMPLATE, OUTPUT)
    doc = Document(str(OUTPUT))

    info_table = doc.tables[0]
    info_table.cell(0, 0).text = "专业 ________    班级 ________    姓名 ________    学号 ________________"
    info_table.cell(1, 0).text = "课程名称 微控制器与嵌入式系统实验              实验名称 实验六 电机控制实验"
    info_table.cell(2, 0).text = "地点 TDX-PITE实验平台 / Proteus仿真    台号 ______    指导教师 ________    日期 ________"

    content_cell = doc.tables[1].cell(0, 0)
    content_cell.text = ""

    add_heading(content_cell, "一、实验内容、目的与要求")
    add_paragraph(
        content_cell,
        "本实验对应《2026微控制器与嵌入式系统实验任务书（适用于自动化、测控2024级）》中的实验六“基于 MCS51 单片机步进电机、直流电机控制实验”。基本题包括《单片机实验指导.pdf》5.1 步进电机实验和 5.2 直流电机 PWM 调速实验；提高题要求设计一个直流电机控制系统，支持按键启动、停止、速度增减，速度等级为 0-4，数码管显示当前状态，上位机实时监控并可发送控制命令，在速度边界继续加减速时进行声光报警并发送 error。"
    )
    add_paragraph(
        content_cell,
        "实验目的包括：理解步进电机四相八拍励磁顺序和端口输出控制方法；理解 PWM 占空比改变直流电机平均驱动电压的原理；掌握定时器中断、按键消抖、串口通信、状态显示和异常报警在同一控制程序中的协调设计。"
    )
    add_paragraph(
        content_cell,
        "实验要求是基本题能够分别产生步进电机相序和直流电机 PWM 波形，提高题能够形成闭环的人机交互控制逻辑。程序需要在 C51 环境下编译通过，并能通过 Keil Debug、Proteus 逻辑分析仪、串口窗口或实验箱硬件观察到明确的运行效果。"
    )

    add_heading(content_cell, "二、实验硬件与软件环境条件")
    add_table(
        content_cell,
        ["类别", "名称", "作用"],
        [
            ["硬件", "TDX-PITE 教学实验系统 / TD-51 单片机平台", "提供 MCS-51 控制核心和实验接口"],
            ["硬件", "四相步进电机及驱动模块", "验证相序输出和方向控制"],
            ["硬件", "直流电机及驱动模块", "验证 PWM 调速效果"],
            ["硬件", "按键、数码管、LED、蜂鸣器", "实现启动停止、速度显示和边界报警"],
            ["硬件", "串口通信模块", "实现上位机监控和远程控制"],
            ["软件", "Keil uVision5 + PK51 C51", "编写、编译、调试 8051 程序"],
            ["软件", "Proteus 8 Professional / 串口调试助手", "观察端口波形、PWM 占空比和串口文本"],
        ],
    )
    add_paragraph(
        content_cell,
        "本实验形成 3 个源程序：EX6_Stepper.c 用于步进电机基本题，EX6_DC_PWM.c 用于直流电机 PWM 基本题，EX6_DC_Control_Advanced.c 用于提高题综合控制。三份程序均已使用 C51 编译器检查，结果为 0 个错误、0 个警告。"
    )

    add_heading(content_cell, "三、实验原理")
    add_heading(content_cell, "1. 步进电机控制原理", level=2)
    add_paragraph(
        content_cell,
        "步进电机通过定子绕组按一定顺序通电，使转子按固定步距角逐步转动。四相八拍方式在单相通电和双相通电之间交替，能够比四拍方式获得更细的步进过程。实验中 P0.0-P0.3 分别控制 A、B、C、D 四相，输出序列为 01H、03H、02H、06H、04H、0CH、08H、09H。按该序列递增输出时电机按一个方向转动，递减输出时方向相反。"
    )
    add_heading(content_cell, "2. 直流电机 PWM 调速原理", level=2)
    add_paragraph(
        content_cell,
        "直流电机转速与电枢平均电压有关。PWM 不是连续改变电压幅值，而是在固定周期内改变高电平持续时间。占空比 D=Ton/(Ton+Toff)，平均电压近似为 Uavg=D x Vcc。占空比越大，电机获得的平均能量越大，转速越高。实验基本题用 P1.7 输出方波，通过改变高低电平延时得到不同占空比。"
    )
    add_heading(content_cell, "3. 提高题综合控制原理", level=2)
    add_paragraph(
        content_cell,
        "提高题使用 Timer0 产生 1 ms 周期中断，在中断中维护 pwm_count，并根据 speed 查 duty_table 决定 P1.7 是否为高电平。PWM_PERIOD=20，因此 PWM 周期约为 20 ms，速度等级 0-4 对应 duty_table 中的 0、5、10、15、20，即 0%、25%、50%、75%、100% 五档占空比。"
    )
    add_paragraph(
        content_cell,
        "主循环负责处理按键和串口命令。按键采用低电平有效并进行软件消抖，P3.2 启动、P3.3 停止、P3.4 加速、P3.5 减速。串口命令 s、p、+、-、0-4 分别对应启动、停止、加速、减速和直接设定速度。速度为 4 时继续加速或速度为 0 时继续减速，会调用 trigger_error()，使报警 LED 和蜂鸣器动作，并向上位机发送 error。"
    )

    add_heading(content_cell, "四、实验线路示图与端口说明")
    add_paragraph(
        content_cell,
        "步进电机实验线路为：P0.0-P0.3 连接步进电机驱动模块的 A、B、C、D 四相控制输入，驱动模块再接步进电机绕组。由于单片机端口不能直接驱动电机绕组，实际硬件必须经过达林顿管、驱动芯片或实验箱内置功率驱动电路。"
    )
    add_paragraph(
        content_cell,
        "直流电机 PWM 线路为：P1.7 输出 PWM 控制信号，经驱动电路控制直流电机通断。提高题中 P1.0 接报警 LED，P1.1 接蜂鸣器，P2 接一位共阴极数码管段码，P3.2-P3.5 接四个独立按键，TXD/RXD 接串口通信模块。停止状态数码管显示 '-'，运行时显示当前速度等级 0-4。"
    )
    add_table(
        content_cell,
        ["端口/变量", "连接或功能", "说明"],
        [
            ["P0.0-P0.3", "步进电机 A-D 相控制", "按八拍励磁表循环输出"],
            ["P1.7", "直流电机 PWM 输出", "占空比由 speed 决定"],
            ["P1.0", "报警 LED", "越界操作时点亮"],
            ["P1.1", "蜂鸣器", "越界操作时发出间歇报警"],
            ["P2", "数码管段码", "显示 '-' 或速度等级"],
            ["P3.2/P3.3", "启动/停止键", "低电平有效"],
            ["P3.4/P3.5", "加速/减速键", "低电平有效"],
            ["TXD/RXD", "上位机串口", "发送 RUN、SPEED、error 并接收命令"],
        ],
    )
    add_paragraph(content_cell, "算法流程图已整理在 EX6/Paper/实验六算法流程图_mermaid.md 中，包含步进电机、直流 PWM 和综合控制提高题三个流程。")

    add_heading(content_cell, "五、关键程序设计与分析")
    add_heading(content_cell, "1. 步进电机基本题", level=2)
    add_paragraph(
        content_cell,
        "EX6_Stepper.c 中 step_table 保存八拍相序。主循环根据 P3.2 的状态决定 index 递增或递减，并将 step_table[index] 输出到 P0 低四位。通过改变 delay_ms() 的延时时间，可以改变每拍间隔，从而改变步进电机转速。"
    )
    add_heading(content_cell, "2. PWM 基本题", level=2)
    add_paragraph(
        content_cell,
        "EX6_DC_PWM.c 中 P1.7 先保持高电平 T_HIGH 个延时单位，再保持低电平 T_LOW 个延时单位。高低电平延时之比决定占空比。该方法结构直观，适合基本题验证 PWM 概念，但主循环完全被延时占用，不适合同时处理按键和串口，因此提高题改用定时器中断产生 PWM。"
    )
    add_heading(content_cell, "3. 提高题定时器与状态机", level=2)
    add_paragraph(
        content_cell,
        "EX6_DC_Control_Advanced.c 使用 Timer0 方式 1，装入 0xFC18。对 12 MHz、12T 8051 来说，机器周期为 1 us，65536-0xFC18=1000，因此每次溢出约为 1 ms。中断中递增 pwm_count 并与 duty_table[speed] 比较，由硬件节拍保证 PWM 周期稳定。"
    )
    add_paragraph(
        content_cell,
        "程序状态由 running、speed、alarm_ms 三个核心变量描述。running 决定电机是否输出 PWM，speed 表示速度等级，alarm_ms 表示报警剩余时间。主循环不断处理按键和串口，定时器中断负责 PWM、报警闪烁和 1 s 状态上报标志 status_due。这样避免长延时阻塞，提高了系统响应性。"
    )
    add_heading(content_cell, "4. 边界报警逻辑", level=2)
    add_paragraph(
        content_cell,
        "speed_up() 在 speed<4 时加 1，否则触发 error；speed_down() 在 speed>0 时减 1，否则触发 error。这种写法把边界检查集中在调速函数内，按键和串口共用同一套规则，避免出现按键能报警而串口不报警的逻辑不一致。"
    )

    add_heading(content_cell, "六、实验调试步骤")
    add_paragraph(content_cell, "（1）在 Keil 中分别建立 3 个工程，每个工程只加入一个源文件，选择 SST89C52 或兼容 8051 器件，编译确认无错误和警告。")
    add_paragraph(content_cell, "（2）步进电机基本题进入 Debug 后打开 Port 0 或 Logic Analyzer，观察 P0 低四位是否按 01、03、02、06、04、0C、08、09 循环。改变 P3.2 后，序列方向应反向。")
    add_paragraph(content_cell, "（3）直流 PWM 基本题在 Logic Analyzer 中添加 P1.7，观察周期性方波。修改 T_HIGH 和 T_LOW 后重新编译，方波高低电平比例应随之改变。")
    add_paragraph(content_cell, "（4）提高题调试时，在 Watch 窗口加入 running、speed、pwm_count、alarm_ms、status_due；在 Logic Analyzer 中观察 P1.7；在 Serial Window #1 中观察 RUN=..., SPEED=... 和 error。")
    add_paragraph(content_cell, "（5）分别测试按键和串口两类控制路径：按启动键或发送 s 后 running=1；按停止键或发送 p 后 running=0；按加减速键或发送 +/- 后 speed 改变；发送 0-4 可直接设定速度等级。")
    add_paragraph(content_cell, "（6）边界测试是提高题验收重点：speed=4 时继续加速应发送 error 并触发声光报警；speed=0 时继续减速也应产生同样报警。若没有报警，应检查 speed_up()、speed_down() 和 trigger_error() 是否被调用。")

    add_heading(content_cell, "七、实验数据记录及结果分析")
    add_table(
        content_cell,
        ["速度等级", "duty_table", "理论占空比", "预期现象"],
        [
            ["0", "0/20", "0%", "运行时 PWM 始终低，电机停止或最低速"],
            ["1", "5/20", "25%", "P1.7 高电平约占四分之一周期"],
            ["2", "10/20", "50%", "P1.7 高低电平基本相等"],
            ["3", "15/20", "75%", "P1.7 高电平约占四分之三周期"],
            ["4", "20/20", "100%", "P1.7 基本保持高电平，最高速"],
        ],
    )
    add_table(
        content_cell,
        ["操作", "状态变化", "串口输出", "数码/报警现象"],
        [
            ["启动键或 s", "running=1", "RUN=1, SPEED=x", "显示当前速度等级"],
            ["停止键或 p", "running=0", "RUN=0, SPEED=x", "数码管显示 '-'，PWM 停止"],
            ["加速键或 +", "speed 加 1", "RUN=x, SPEED+1", "PWM 占空比增大"],
            ["减速键或 -", "speed 减 1", "RUN=x, SPEED-1", "PWM 占空比减小"],
            ["speed=4 继续加速", "speed 保持 4", "error", "LED 亮，蜂鸣器报警"],
            ["speed=0 继续减速", "speed 保持 0", "error", "LED 亮，蜂鸣器报警"],
        ],
    )
    add_paragraph(
        content_cell,
        "结果分析表明，基本题和提高题的控制对象虽然都是电机，但程序组织方式不同。基本 PWM 题用延时即可观察波形；提高题需要同时处理按键、串口和报警，必须把 PWM 放入定时器中断，主循环只处理低频的人机交互事件。"
    )
    add_paragraph(
        content_cell,
        "Timer0 1 ms 节拍使 PWM 周期为 20 ms，即频率约为 50 Hz。该频率适合在仿真和逻辑分析仪中观察占空比；真实电机控制时可根据驱动电路和电机特性适当提高 PWM 频率，以减小机械抖动和可闻噪声。"
    )

    add_heading(content_cell, "八、误差分析与注意事项")
    add_paragraph(content_cell, "步进电机控制中，相序错误会导致电机抖动、反转或不转；延时过短可能使电机来不及响应而失步。调试时应先用较慢节拍确认相序正确，再逐步减小延时。")
    add_paragraph(content_cell, "直流电机 PWM 调速中，软件延时方式受编译优化和主循环负载影响，波形稳定性不如定时器中断方式。提高题使用 Timer0 后，占空比由中断节拍决定，稳定性更好。")
    add_paragraph(content_cell, "声光报警属于异常状态提示，应与速度状态分离。程序中报警持续时间由 alarm_ms 控制，报警期间不改变 speed，这样可以避免误操作导致状态不可追踪。")
    add_paragraph(content_cell, "如果串口输出乱码，需要检查晶振频率、TH1、SMOD 和上位机波特率；如果 P1.7 无波形，需要确认 EA、ET0、TR0 已置位，并检查 Timer0 是否被其他初始化语句覆盖。")

    add_heading(content_cell, "九、实验总结与心得")
    add_paragraph(
        content_cell,
        "实验六从执行机构控制角度综合训练了单片机端口输出、定时器中断和人机交互设计。步进电机强调离散相序和节拍，直流电机强调 PWM 占空比和平均电压，提高题则强调多个功能模块在同一程序中的协调。"
    )
    add_paragraph(
        content_cell,
        "本实验最重要的工程经验是：当系统需要同时处理实时输出和用户输入时，不能依赖长时间阻塞延时。把 PWM 和周期状态维护交给定时器中断，把按键、串口和显示放在主循环中，可以让程序结构更清晰，也更接近真实嵌入式控制系统的设计方式。"
    )

    doc.save(str(OUTPUT))


if __name__ == "__main__":
    build_report()
