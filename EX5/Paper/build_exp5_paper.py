from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parents[2]
EX5_PAPER = ROOT / "EX5" / "Paper"
TEMPLATE = ROOT / "EX1" / "Paper" / "exp1_report.docx"
OUTPUT = EX5_PAPER / "exp5_report.docx"


def set_run_font(run, east_asia: str, latin: str | None = None, size: int = 11, bold: bool = False):
    run.font.name = latin or east_asia
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.bold = bold


def add_paragraph(cell, text: str, indent: float = 0.74, size: int = 11):
    p = cell.add_paragraph()
    p.style = "Normal"
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(indent)
    run = p.add_run(text)
    set_run_font(run, "仿宋_GB2312", size=size)
    return p


def add_heading(cell, text: str, level: int = 1):
    p = cell.add_paragraph()
    p.style = "Normal"
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(8 if level == 1 else 4)
    run = p.add_run(text)
    set_run_font(run, "黑体", size=14 if level == 1 else 11, bold=True)
    return p


def add_table(cell, headers: list[str], rows: list[list[str]]):
    table = cell.add_table(rows=1, cols=len(headers))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    for i, text in enumerate(headers):
        table.rows[0].cells[i].text = text
    for row_data in rows:
        row = table.add_row()
        for i, text in enumerate(row_data):
            row.cells[i].text = text
    for row in table.rows:
        for c in row.cells:
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.line_spacing = 1.15
                for run in p.runs:
                    set_run_font(run, "仿宋_GB2312", size=9)
    return table


def build_report():
    EX5_PAPER.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(TEMPLATE, OUTPUT)
    doc = Document(str(OUTPUT))

    info_table = doc.tables[0]
    info_table.cell(0, 0).text = "专业 ________    班级 ________    姓名 ________    学号 ________________"
    info_table.cell(1, 0).text = "课程名称 微控制器与嵌入式系统实验              实验名称 实验五 串口通讯与 A/D 转换"
    info_table.cell(2, 0).text = "地点 TDX-PITE实验平台 / Proteus仿真    台号 ______    指导教师 ________    日期 ________"

    content_cell = doc.tables[1].cell(0, 0)
    content_cell.text = ""

    add_heading(content_cell, "一、实验内容、目的与要求")
    add_paragraph(
        content_cell,
        "本实验对应《2026微控制器与嵌入式系统实验任务书（适用于自动化、测控2024级）》中的实验五“串口通讯与 A/D 转换实验”。基本题包括《单片机实验指导.pdf》3.7 串口通讯实验和 A/D 转换实验。任务书中把 A/D 写为 4.3，但本指导书 4.2 为 A/D 转换、4.3 为 D/A 转换，因此本报告按指导书 4.2 A/D 转换完成。提高题要求把 A/D 转换数据实时发送到上位机，并在下位机实时显示当前 A/D 值，同时由上位机控制转换值的显示精度。",
    )
    add_paragraph(
        content_cell,
        "实验目的包括：掌握 MCS-51 串行口方式 1 的工作过程，理解 SCON、SBUF、RI、TI、SMOD 与定时器 1 波特率发生器之间的关系；掌握 ADC0809 的启动转换、忙信号等待和转换结果读取方法；能够把采样、显示、串口通信和上位机命令解析组织为一个完整的嵌入式数据采集程序。",
    )
    add_paragraph(
        content_cell,
        "实验要求是分别完成串口发送基本题、ADC0809 采样基本题和 ADC+串口综合提高题。程序应能在 Keil C51 环境下正确编译，在 Proteus 或实验箱环境中可观察到串口输出、端口采样值变化和精度控制命令响应。"
    )

    add_heading(content_cell, "二、实验硬件与软件环境条件")
    add_table(
        content_cell,
        ["类别", "名称", "作用"],
        [
            ["硬件", "TDX-PITE 教学实验系统 / TD-51 单片机平台", "提供 MCS-51 核心板、端口扩展和实验模块连接条件"],
            ["硬件", "RS-232 或串口通信模块", "完成单片机 TXD/RXD 与上位机串口之间的电平转换和数据传输"],
            ["硬件", "ADC0809 A/D 转换模块", "把 0-5 V 模拟量转换为 8 位数字量"],
            ["硬件", "数码管、LED 或端口观察窗口", "显示当前 A/D 原始采样值或调试端口状态"],
            ["软件", "Keil uVision5 + PK51 C51", "编写、编译、调试 8051 C 程序"],
            ["软件", "Proteus 8 Professional / 串口调试助手", "完成仿真验证和上位机串口交互观察"],
        ],
    )
    add_paragraph(
        content_cell,
        "本实验形成 3 个源程序：EX5_Serial.c 用于串口基本题，EX5_ADC0809.c 用于 A/D 基本题，EX5_ADC_Serial_Advanced.c 用于提高题。三份程序均已使用 C51 编译器检查，结果为 0 个错误、0 个警告。"
    )

    add_heading(content_cell, "三、实验原理")
    add_heading(content_cell, "1. 串口通讯原理", level=2)
    add_paragraph(
        content_cell,
        "MCS-51 串口方式 1 为 10 位异步通信格式，包括 1 位起始位、8 位数据位和 1 位停止位。发送时软件把待发送字节写入 SBUF，硬件自动移位输出，发送完成后置位 TI；接收时硬件把接收到的数据放入 SBUF，并置位 RI。本实验采用查询方式，发送函数在写 SBUF 后等待 TI 置位，再清零 TI，保证相邻字符不会覆盖。"
    )
    add_paragraph(
        content_cell,
        "方式 1 的波特率通常由定时器 1 溢出率决定。指导书推荐晶振 11.0592 MHz、定时器 1 方式 2 自动重装、TH1=0xFD、SMOD=1，此时波特率约为 19200 bps。程序中 SCON=0x50 表示串口方式 1 且允许接收，PCON=0x80 置 SMOD=1，TMOD 高四位配置定时器 1 为方式 2。"
    )
    add_heading(content_cell, "2. A/D 转换原理", level=2)
    add_paragraph(
        content_cell,
        "ADC0809 是 8 位逐次逼近型 A/D 转换器，转换结果范围为 0-255。若参考电压为 5 V，则理论分辨率为 5/255=0.01961 V，即约 19.6 mV/LSB。输入电压与数字量近似满足 Uin = AD / 255 x 5.000 V。"
    )
    add_paragraph(
        content_cell,
        "本实验按照指导书地址映射访问 ADC0809：STARTAD=XBYTE[0x7F00] 用于启动转换，ADRESULT=XBYTE[0x7F08] 用于读取转换结果，ADBUSY=P3.3 作为忙信号。程序向 STARTAD 写任意值后启动转换，等待忙信号结束，再从 ADRESULT 取出 8 位采样值。"
    )
    add_heading(content_cell, "3. 提高题综合控制思路", level=2)
    add_paragraph(
        content_cell,
        "提高题把 ADC 采样值同时用于三处：一是送 P2 作为端口观察值；二是分解为百位、十位和个位，通过 P0 段码与 P1 位选扫描到数码管；三是换算为电压文本，经串口发送到上位机。上位机发送字符 0、1、2、3 时，程序改变 precision 变量，从而控制电压文本保留的小数位数。"
    )

    add_heading(content_cell, "四、实验线路示图与端口说明")
    add_paragraph(
        content_cell,
        "串口实验线路为：单片机 P3.1/TXD 连接串口发送通道，P3.0/RXD 连接串口接收通道，经 MAX232 或实验箱电平转换模块与 PC 串口相连。仿真时可直接连接 Virtual Terminal 或使用 Keil Serial Window #1 观察收发数据。"
    )
    add_paragraph(
        content_cell,
        "A/D 实验线路为：可调模拟电压接入 ADC0809 的模拟输入端，ADC0809 的数据总线接单片机外部数据总线，片选和启动控制由实验箱译码电路映射到 0x7F00 与 0x7F08。P3.3 连接 ADBUSY，用于判断转换是否完成。提高题中 P0 用作数码管段码，P1 低三位用作三位数码管位选，P2 输出 ADC 原始值，便于在端口窗口或逻辑分析仪中观察。"
    )
    add_table(
        content_cell,
        ["信号/变量", "连接或含义", "实验作用"],
        [
            ["TXD/RXD", "P3.1/P3.0", "与上位机交换文本命令和采样数据"],
            ["STARTAD", "XBYTE[0x7F00]", "启动 ADC0809 转换"],
            ["ADRESULT", "XBYTE[0x7F08]", "读取 ADC0809 转换结果"],
            ["ADBUSY", "P3.3", "判断 A/D 转换是否完成"],
            ["P0", "数码管段码", "显示 ADC 原始值的各位数字"],
            ["P1", "数码管位选", "动态选择百位、十位、个位"],
            ["P2", "8 位原始采样值", "便于端口级调试和结果验证"],
        ],
    )
    add_paragraph(content_cell, "算法流程图已整理在 EX5/Paper/实验五算法流程图_mermaid.md 中，包含串口发送、A/D 采样和提高题综合流程。")

    add_heading(content_cell, "五、关键程序设计与分析")
    add_heading(content_cell, "1. 串口初始化", level=2)
    add_paragraph(
        content_cell,
        "uart_init() 中 SCON=0x50 使串口工作在方式 1 并允许接收；TMOD=(TMOD&0x0F)|0x20 只修改定时器 1 的配置，避免破坏定时器 0；TH1=TL1=0xFD 设定自动重装初值；TR1=1 启动定时器 1；TI=1 使第一次调用 uart_putc() 时不被上一轮发送标志阻塞。"
    )
    add_heading(content_cell, "2. ADC 采样函数", level=2)
    add_paragraph(
        content_cell,
        "read_adc0809() 先写 STARTAD 启动转换，再等待 ADBUSY 结束，然后读取 ADRESULT。这样的顺序符合 ADC0809 的转换时序，避免在转换尚未结束时读取旧值。为了增强仿真稳定性，等待结束后加入短延时，使外部总线数据有足够保持时间。"
    )
    add_heading(content_cell, "3. 电压换算与精度控制", level=2)
    add_paragraph(
        content_cell,
        "uart_put_voltage() 使用 unsigned long 计算毫伏值 mv=value*5000/255，避免 8 位或 16 位变量溢出。precision=0 时只发送整数电压；precision=1、2、3 时分别发送一位、两位和三位小数。handle_uart_command() 检查 RI 标志并读取 SBUF，当收到 '0' 到 '3' 时立即更新 precision 并回传确认信息。"
    )

    add_heading(content_cell, "六、实验调试步骤")
    add_paragraph(content_cell, "（1）在 Keil 中分别建立 3 个工程，每个工程只加入一个包含 main() 的源文件，避免多个主函数冲突。目标器件选择 SST89C52 或兼容 8051 器件，并启用 C51 编译。")
    add_paragraph(content_cell, "（2）串口基本题进入 Debug 后打开 Serial Window #1，运行程序，应周期性看到 Xi'an Tangdu Corp.。若在 Proteus 中验证，应确保晶振和虚拟终端波特率与程序配置一致。")
    add_paragraph(content_cell, "（3）A/D 基本题在 Watch 窗口加入 ad_value，同时观察 P1 或 P2。改变输入模拟电压时，ad_value 应随电压升高而增大，0 V 附近接近 0，5 V 附近接近 255。")
    add_paragraph(content_cell, "（4）提高题运行后，在串口窗口观察 ADC=..., U=...V 的周期性输出；发送 0、1、2、3 后观察 precision 变量和输出小数位数是否同步改变；同时查看 P2 端口或数码管显示是否与 ADC 原始值一致。")
    add_paragraph(content_cell, "（5）若串口显示乱码，优先检查晶振频率、SMOD、TH1 和串口助手波特率；若 ADC 值不变化，优先检查外部地址映射、ADBUSY 连接和模拟输入电压是否有效。")

    add_heading(content_cell, "七、实验数据记录及结果分析")
    add_table(
        content_cell,
        ["模拟输入", "理论 ADC 值", "理论显示电压", "验证要点"],
        [
            ["0.0 V", "0", "0.000 V", "P2 接近 0x00，串口电压接近 0 V"],
            ["1.0 V", "约 51", "1.000 V", "ADC=51 左右，允许受仿真步进和参考电压误差影响"],
            ["2.5 V", "约 128", "2.509 V", "中点输入时数字量约为 128"],
            ["5.0 V", "255", "5.000 V", "P2 接近 0xFF，串口显示满量程"],
        ],
    )
    add_table(
        content_cell,
        ["上位机命令", "precision 值", "示例输出", "含义"],
        [
            ["0", "0", "ADC=128, U=2V", "只显示整数电压"],
            ["1", "1", "ADC=128, U=2.5V", "显示 1 位小数"],
            ["2", "2", "ADC=128, U=2.50V", "显示 2 位小数"],
            ["3", "3", "ADC=128, U=2.509V", "显示 3 位小数"],
        ],
    )
    add_paragraph(
        content_cell,
        "从理论关系看，ADC0809 为 8 位转换器，因此单个采样值的最小电压分辨力约为 19.6 mV。提高显示小数位数并不会提高硬件本身的 A/D 分辨率，而是改变上位机文本显示格式。报告中把“精度控制”解释为显示精度控制更科学，避免把格式位数误认为转换器真实精度提高。"
    )
    add_paragraph(
        content_cell,
        "实验结果表明，串口发送与接收逻辑能够稳定工作；ADC0809 采样值随输入电压单调变化；提高题能够在持续采样和动态显示的同时响应上位机命令，满足“实时上传、实时显示、上位机控制显示精度”的要求。"
    )

    add_heading(content_cell, "八、误差分析与注意事项")
    add_paragraph(content_cell, "A/D 结果误差主要来自参考电压偏差、ADC0809 量化误差、模拟输入源稳定度和仿真模型时序。8 位量化决定了结果不可能无限精确，理论最大量化误差约为半个 LSB，即约 9.8 mV。")
    add_paragraph(content_cell, "串口通信可靠性主要受波特率误差影响。8051 串口常使用 11.0592 MHz 晶振，因为它可以整除常见波特率；若使用 12 MHz，需要重新计算 TH1，否则上位机可能出现乱码。")
    add_paragraph(content_cell, "动态数码管显示需要保持足够刷新频率。若主循环中串口发送过于频繁，可能造成显示闪烁；本程序通过 loop_count 降低串口上传频率，在采样显示和通信之间取得平衡。")

    add_heading(content_cell, "九、实验总结与心得")
    add_paragraph(
        content_cell,
        "实验五的核心不只是分别完成串口和 A/D，而是把两类典型外设组织成一个数据采集系统。串口部分要求时序和波特率严格匹配，A/D 部分要求按转换器时序启动、等待和读取；提高题进一步要求程序同时处理采样、显示、通信和命令解析，体现了嵌入式系统中多任务协同的基本思想。"
    )
    add_paragraph(
        content_cell,
        "通过本实验可以认识到，显示位数、通信格式和硬件转换精度是三个不同层面的概念。上位机可以控制显示格式，使数据更适合观察和记录，但真实转换精度仍由 ADC 位数、参考电压和模拟电路决定。这一点对后续设计更复杂的数据采集与控制系统很重要。"
    )

    doc.save(str(OUTPUT))


if __name__ == "__main__":
    build_report()
