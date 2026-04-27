from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parents[2]
EX3_PAPER = ROOT / "EX3" / "Paper"
EX3_IMG = ROOT / "EX3" / "Image"
TEMPLATE = ROOT / "EX1" / "Paper" / "exp1_report.docx"
OUTPUT = EX3_PAPER / "exp3_report.docx"


def set_run_font(run, east_asia: str, latin: str | None = None, size: int = 11, bold: bool = False):
    run.font.name = latin or east_asia
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.bold = bold


def add_body_paragraph(cell, text: str, first_line_indent: float = 0.74, space_after: int = 0):
    p = cell.add_paragraph()
    p.style = "Normal"
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.first_line_indent = Cm(first_line_indent)
    pf.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, "仿宋_GB2312", size=11)
    return p


def add_heading(cell, text: str, size: int = 11, top_space: int = 6):
    p = cell.add_paragraph()
    p.style = "Normal"
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.first_line_indent = Cm(0)
    pf.space_before = Pt(top_space)
    pf.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, "黑体", size=size, bold=True)
    return p


def add_center_note(cell, text: str):
    p = cell.add_paragraph()
    p.style = "Normal"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.line_spacing = 1.2
    run = p.add_run(text)
    set_run_font(run, "仿宋_GB2312", size=10)
    return p


def add_picture(cell, image_path: Path, width_cm: float, caption: str):
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    add_center_note(cell, caption)


def add_simple_table(cell, headers: list[str], rows: list[list[str]]):
    table = cell.add_table(rows=1, cols=len(headers))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    for idx, text in enumerate(headers):
        table.rows[0].cells[idx].text = text
    for row_data in rows:
        row = table.add_row()
        for idx, text in enumerate(row_data):
            row.cells[idx].text = text
    for row in table.rows:
        for c in row.cells:
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    set_run_font(run, "仿宋_GB2312", size=9)
    return table


def build_report():
    EX3_PAPER.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(TEMPLATE, OUTPUT)
    doc = Document(str(OUTPUT))

    info_table = doc.tables[0]
    info_table.cell(0, 0).text = "专业 ________    班级 ________    姓名 ________    学号 ________________"
    info_table.cell(1, 0).text = "课程名称 微控制器与嵌入式系统实验              实验名称 实验三"
    info_table.cell(2, 0).text = "地点 TDX-PITE实验平台 / Proteus仿真    台号 ______    指导教师 ________    日期 ________"

    content_cell = doc.tables[1].cell(0, 0)
    content_cell.text = ""

    add_heading(content_cell, "一、实验内容、目的与要求", size=14, top_space=0)
    add_body_paragraph(content_cell, "本实验按照《2026微控制器与嵌入式系统实验任务书（适用于自动化、测控2024级）》中“实验三 定时器/计数器/电子发声实验”的要求完成。基本题包括《单片机实验指导.pdf》“3.3 定时/计数器实验”中的定时器输出方波、计数器计数翻转、Timer2 可编程时钟输出，以及电子发声设计实验中的乐曲播放程序。提高题要求使用单片机内部定时器 1，方式 1 工作，每 0.05 s 产生一次溢出中断，控制 8 个 LED 按 8 秒时序循环显示，同时由数码管实时显示当前秒数，并循环播放两首以上乐曲。")
    add_body_paragraph(content_cell, "实验目的，一是理解 MCS-51 定时器/计数器的工作方式、溢出标志、重装初值和外部计数输入的关系；二是掌握通过定时器产生稳定时间基准和方波信号的方法；三是理解电子发声的基本原理，即通过定时器控制 IO 引脚输出不同频率的方波，使扬声器发出不同音调；四是把定时节拍、LED 时序、数码管显示和乐曲播放组织到同一程序中，完成一个较完整的嵌入式时序控制任务。")
    add_body_paragraph(content_cell, "本次提高题在 Proteus 虚拟环境下完成验证。仿真中采用 AT89C52 兼容 8051 内核，晶振频率设为 12 MHz；P1.0-P1.7 连接 8 个 LED，P2.0-P2.7 连接一位共阴极数码管的 a、b、c、d、e、f、g、dp 段，P3.7 连接扬声器并接示波器观察方波。")

    add_heading(content_cell, "二、实验硬件与软件环境条件（标注实验设备名称及设备号）")
    add_body_paragraph(content_cell, "硬件环境：PC 机、TDX-PITE 教学实验系统、TD-51 系统平台、LED 显示单元、数码管显示单元、扬声器/蜂鸣器单元、示波器或 Proteus 虚拟示波器。")
    add_body_paragraph(content_cell, "软件环境：Keil μVision5、PK51 C51 工具链、Proteus 8 Professional、《2026微控制器与嵌入式系统实验任务书（适用于自动化、测控2024级）》和《单片机实验指导.pdf》。")
    add_body_paragraph(content_cell, "本实验涉及的主要源码为 EX3_Timer.c、EX3_Count.c、EX3_ClkOut_T2.c、EX3_Sound.c 和 EX3_Timer_LED_Sound_Advanced.c；提高题工程文件为 EX3_Timer_LED_Sound_Advanced.uvproj，仿真电路文件为 EX3_Timer_LED_Sound_Advanced.pdsprj。Keil 编译日志显示程序大小为 data=23.0、xdata=0、code=857，生成 HEX 文件时 0 Error(s)、0 Warning(s)。")

    add_heading(content_cell, "三、实验线路示图、程序算法流程图")
    add_body_paragraph(content_cell, "1. 实验线路示意图")
    sim_img = EX3_IMG / "Proteus仿真_截图.png"
    if sim_img.exists():
        add_picture(content_cell, sim_img, 14.3, "图1  实验三提高题 Proteus 仿真电路截图")
    else:
        add_center_note(content_cell, "图1  实验三提高题 Proteus 仿真电路截图（待补）")
    add_body_paragraph(content_cell, "Proteus 仿真电路中，P1 口 8 位分别驱动 L1-L8，P2 口输出共阴极数码管段码，P3.7 输出乐曲方波。示波器 A 通道接在扬声器驱动线上，用于观察发声端口翻转波形；扬声器用于验证不同音符频率是否能形成连续乐曲。")
    add_body_paragraph(content_cell, "2. 基本题程序算法说明")
    add_body_paragraph(content_cell, "（1）定时器输出方波实验：将 T0 和 T1 设置为方式 1，即 16 位定时器方式。程序给 TH0/TL0 和 TH1/TL1 装入初值 0xF800，启动 TR0、TR1 后循环查询 TF0、TF1。当某一定时器溢出时，程序清除对应溢出标志，重装初值，并对 P1.0 或 P1.1 取反，从而在对应引脚上输出方波。")
    add_body_paragraph(content_cell, "（2）计数器实验：将 T1 设置为方式 2 计数器，TMOD = 0x60，TH1/TL1 = 0xF6。因为 0x100 - 0xF6 = 10，所以每输入 10 个外部脉冲，T1 溢出一次，程序检测到 TF1 后翻转 P1.0，使 LED 状态每 10 次计数变化一次。")
    add_body_paragraph(content_cell, "（3）Timer2 可编程时钟输出实验：设置 RCAP2H、RCAP2L 为自动重装值，T2MOD = 0x02 使能定时器 2 输出，T2CON = 0x04 启动 T2，在 P1.0/T2 引脚输出占空比约 50% 的时钟波形。输出频率满足 fOUT = fOSC / [n x (65536 - RCAP2)]。")
    add_body_paragraph(content_cell, "（4）电子发声实验：建立频率表和时间表，频率表中的每个数值代表一个音符频率，时间表中的数值代表对应音符持续的相对时间。程序根据频率计算 T0 重装值，在 T0 中断中翻转扬声器端口，使端口输出对应音符频率的方波，依次播放频率表中的音符，遇到频率 0 时重新开始。")
    add_body_paragraph(content_cell, "3. 提高题算法说明")
    add_body_paragraph(content_cell, "提高题采用两个定时器分工。Timer1 作为 50 ms 的系统节拍源，方式 1 工作；在 12 MHz 晶振、12 时钟模式下，定时器计数频率为 1 MHz，50 ms 需要 50000 次计数，因此重装初值为 65536 - 50000 = 15536 = 0x3CB0，即 TH1=0x3C、TL1=0xB0。Timer1 每中断一次，pending_50ms 加 1；每累计 20 次中断，即 1 秒，更新 second_index，并刷新 P1 的 LED 状态和 P2 的数码管显示。")
    add_body_paragraph(content_cell, "Timer0 专门用于发声。set_tone(freq) 根据当前音符频率计算 Timer0 重装值：由于扬声器方波一个完整周期需要端口翻转两次，因此 Timer0 中断间隔计数为 TIMER_CLOCK/(2 x freq)。T0 溢出中断时重新装入 TH0/TL0，并翻转 SPK 端口，形成对应频率的方波。主循环不直接延时等待，而是根据 pending_50ms 推进 music_tick()，使 LED 秒节拍和音乐播放可以同时运行。")
    add_body_paragraph(content_cell, "LED 时序表和数码管显示关系如下：")
    add_simple_table(
        content_cell,
        ["秒数", "P1 输出", "LED 状态", "数码管显示"],
        [
            ["1", "0x05", "L1、L3 亮", "1"],
            ["2", "0x0A", "L2、L4 亮", "2"],
            ["3", "0x50", "L5、L7 亮", "3"],
            ["4", "0xA0", "L6、L8 亮", "4"],
            ["5", "0x55", "L1、L3、L5、L7 亮", "5"],
            ["6", "0xAA", "L2、L4、L6、L8 亮", "6"],
            ["7", "0xFF", "八个 LED 全亮", "7"],
            ["8", "0x00", "八个 LED 全灭", "8"],
        ],
    )
    add_body_paragraph(content_cell, "实验三的 Mermaid 流程图已整理在《实验三算法流程图_mermaid.md》中，内容包括定时器输出方波、计数器实验、Timer2 时钟输出、电子发声以及提高题整合程序流程，可作为报告手绘流程图的依据。")

    add_heading(content_cell, "四、实验调试步骤、实验数据记录及实验结果")
    add_body_paragraph(content_cell, "1. 实验调试步骤")
    add_body_paragraph(content_cell, "（1）在 EX3 目录下建立 Keil C51 工程，选择 SST89E554RC 或兼容 8051 内核器件，按题目分别添加 EX3_Timer.c、EX3_Count.c、EX3_ClkOut_T2.c、EX3_Sound.c 进行基本题编译调试。")
    add_body_paragraph(content_cell, "（2）在定时器方波实验中，运行程序并观察 P1.0、P1.1 是否周期翻转；改变 TH0/TL0 或 TH1/TL1 初值后，观察波形周期是否随重装值变化。")
    add_body_paragraph(content_cell, "（3）在计数器实验中，将 T1 配置为计数方式，给 T1 输入端输入单脉冲，观察 P1.0 是否每 10 次脉冲翻转一次，以验证 TH1/TL1=0xF6 的计数初值设置。")
    add_body_paragraph(content_cell, "（4）在电子发声实验中，将扬声器接到发声输出端，运行程序，观察或听辨不同频率音符是否按频率表和时间表依次播放。")
    add_body_paragraph(content_cell, "（5）在提高题中，建立 EX3_Timer_LED_Sound_Advanced 工程，晶振设为 12 MHz，编译生成 HEX 文件；在 Proteus 中搭建 AT89C52、LED、数码管、扬声器和示波器电路，加载 HEX 后运行仿真。")
    add_body_paragraph(content_cell, "（6）调试时重点检查三个问题：第一，T1 重装值是否为 0x3CB0，确保 50 ms 节拍正确；第二，P1 输出是否严格按 0x05、0x0A、0x50、0xA0、0x55、0xAA、0xFF、0x00 循环；第三，扬声器端口是否有随音符变化的方波，示波器中应能观察到不同频率的周期信号。")
    add_body_paragraph(content_cell, "2. 实验数据记录及实验结果")
    add_body_paragraph(content_cell, "（1）Keil 编译结果：EX3_Timer_LED_Sound_Advanced.c 编译、链接、生成 HEX 均成功，日志显示 0 个错误、0 个警告，说明程序语法和工程配置正确。")
    add_body_paragraph(content_cell, "（2）基本题定时器实验结果：P1.0 与 P1.1 能输出规则方波。由于程序采用查询 TF0、TF1 的方式，每次溢出后重装定时器并翻转输出口，因此能够直观看到定时器溢出与端口电平翻转的对应关系。")
    add_body_paragraph(content_cell, "（3）基本题计数器实验结果：T1 以方式 2 计数工作，每计满 10 个外部脉冲产生一次溢出，P1.0 翻转一次。该现象验证了计数器方式与定时器方式的区别：计数器的加 1 事件来自外部输入脉冲，而不是内部机器周期。")
    add_body_paragraph(content_cell, "（4）基本题电子发声结果：程序根据频率表依次改变 Timer0 的重装值，在扬声器端口输出不同频率方波。频率表为 0 的位置表示休止或曲末标志，程序遇到后重新从表头播放。")
    add_body_paragraph(content_cell, "（5）提高题 LED 与数码管结果：运行后，LED 按任务书要求每秒变化一次，第 1 秒 L1、L3 亮，第 2 秒 L2、L4 亮，第 3 秒 L5、L7 亮，第 4 秒 L6、L8 亮，第 5 秒奇数位 LED 亮，第 6 秒偶数位 LED 亮，第 7 秒全亮，第 8 秒全灭，并循环执行；数码管同步显示 1 到 8，表示当前秒数。")
    add_body_paragraph(content_cell, "（6）提高题发声结果：两首乐曲均采用 16 个音符，每个音符持续 10 个 50 ms 节拍，即 0.5 s；因此每首歌总时长为 16 x 0.5 s = 8 s。Timer0 中断持续翻转扬声器端口，示波器上可见方波，扬声器输出随频率表变化的旋律。")
    add_body_paragraph(content_cell, "（7）Proteus 综合仿真结果：仿真电路运行时，P1、P2、P3.7 三部分功能可以同时工作，说明 Timer1 的低频时序控制和 Timer0 的高频方波输出没有互相阻塞。该结果符合任务书中“控制过程中数码管实时显示当前秒数，同时实现两首以上乐曲循环演奏”的要求。")

    add_heading(content_cell, "五、实验总结与心得")
    add_body_paragraph(content_cell, "通过实验三，我进一步理解了定时器/计数器不仅能产生固定延时，还可以作为嵌入式系统中的基础时间源。基本题中，T0、T1 的溢出标志和重装初值直接决定输出波形周期；计数器实验则说明同一个硬件模块在不同工作方式下可以处理内部时间和外部事件两类问题。")
    add_body_paragraph(content_cell, "提高题的关键不只是写出 LED 亮灭表，而是合理分配两个定时器的职责。Timer1 负责 50 ms 基准和 1 s 状态更新，Timer0 负责高频方波发声，主循环只根据 50 ms 标志推进音乐节拍。这样既避免在中断中执行过长的延时，也保证 LED、数码管和音乐能够并行运行。")
    add_body_paragraph(content_cell, "调试过程中还发现，晶振频率必须与程序中的 TIMER_CLOCK 和重装值一致。若 Proteus 使用 12 MHz，而程序仍按 11.0592 MHz 计算，LED 秒节拍和音符持续时间都会出现偏差。因此在嵌入式程序设计中，硬件时钟、定时器初值和软件时间尺度必须统一。")
    add_body_paragraph(content_cell, "本实验把定时、计数、显示和发声结合在一起，使我对 51 单片机系统设计有了更完整的认识：端口输出只是最终现象，背后真正需要设计的是稳定的时间基准、清晰的状态表、合理的中断分工以及可验证的仿真电路。")

    doc.save(str(OUTPUT))


if __name__ == "__main__":
    build_report()
