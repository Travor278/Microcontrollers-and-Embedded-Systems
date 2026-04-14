from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parents[2]
EX2_PAPER = ROOT / "EX2" / "Paper"
EX2_IMG = ROOT / "EX2" / "Image"
TEMPLATE = ROOT / "EX1" / "Paper" / "exp1_report.docx"
OUTPUT = EX2_PAPER / "exp2_report.docx"


def set_run_font(run, east_asia: str, latin: str | None = None, size: int = 11, bold: bool = False):
    run.font.name = latin or east_asia
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.bold = bold


def add_body_paragraph(cell, text: str, first_line_indent: float = 0.74, space_after: int = 0):
    p = cell.add_paragraph()
    p.style = "Normal"
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.first_line_indent = Cm(first_line_indent)
    pf.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, "仿宋_GB2312", size=11)
    return p


def add_heading(cell, text: str, size: int = 11, top_space: int = 6):
    p = cell.add_paragraph()
    p.style = "Normal"
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.first_line_indent = Cm(0)
    pf.space_before = Pt(top_space)
    pf.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, "黑体", size=size, bold=True)
    return p


def add_center_note(cell, text: str):
    p = cell.add_paragraph()
    p.style = "Normal"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.line_spacing = 1.2
    run = p.add_run(text)
    set_run_font(run, "仿宋_GB2312", size=10)
    return p


def add_picture(cell, image_path: Path, width_cm: float, caption: str):
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    add_center_note(cell, caption)


def build_report():
    EX2_PAPER.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(TEMPLATE, OUTPUT)
    doc = Document(str(OUTPUT))

    info_table = doc.tables[0]
    info_table.cell(0, 0).text = "专业 ________    班级 ________    姓名 ________    学号 ________________"
    info_table.cell(1, 0).text = "课程名称 微控制器与嵌入式系统实验              实验名称 实验二"
    info_table.cell(2, 0).text = "地点 TDX-PITE实验平台 / Proteus仿真    台号 ______    指导教师 ________    日期 ________"

    content_cell = doc.tables[1].cell(0, 0)
    content_cell.text = ""

    add_heading(content_cell, "一、实验内容、目的与要求", size=14, top_space=0)
    add_body_paragraph(content_cell, "本实验按照《2026微控制器与嵌入式系统实验任务书（适用于自动化、测控2024级）》中“实验二 中断系统及 Proteus 仿真”的要求完成。实验内容包括指导书“3.2 中断系统实验”中的定时器中断输出方波实验、外部中断控制 LED 显示实验，以及提高题“带急救车优先控制的交通灯系统”设计与验证。")
    add_body_paragraph(content_cell, "实验的目的，一是理解 MCS-51 单片机中断系统的结构、优先响应方式和中断服务程序设计方法；二是掌握在 Keil μVision 中通过寄存器、Port 窗口、Watch 和 Logic Analyzer 验证中断现象的方法；三是掌握在 Proteus 中搭建 8051 最小系统、加载 HEX 文件并完成虚拟联动仿真的方法；四是通过提高题训练状态机建模能力，把中断请求、时间节拍、端口输出和状态恢复组织成完整控制系统。")
    add_body_paragraph(content_cell, "从实验要求来看，基础题强调‘会配置中断、会看端口结果、会验证现象’，提高题则进一步强调‘在异常事件插入后仍能保持系统行为合理’。其中急救车优先控制不仅要求全红保护，还要求能够在结束后恢复到中断前状态，而不是简单回到初始状态。")
    add_body_paragraph(content_cell, "本次实验以 Keil 仿真与 Proteus 虚拟电路验证相结合的方式完成。基础题主要通过 Keil 中断仿真和 Logic Analyzer 观察波形与端口变化，提高题则在 Keil 中先验证状态机和中断变量，再在 Proteus 中搭建交通灯电路，观察正常轮转、急救车请求、全红保持和恢复原状态等现象。")

    add_heading(content_cell, "二、实验硬件与软件环境条件（标注实验设备名称及设备号）")
    add_body_paragraph(content_cell, "硬件环境：PC 机、TDX-PITE 教学实验系统、TD-51 系统平台、LED 显示与按键单元、联机调试接口、交通灯仿真电路。")
    add_body_paragraph(content_cell, "软件环境：Keil μVision4、Proteus 8 Professional、实验任务书《2026微控制器与嵌入式系统实验任务书（适用于自动化、测控2024级）》和《单片机实验指导.pdf》。")
    add_body_paragraph(content_cell, "本实验涉及的主要源码为 EX2_Int1.c、EX2_INT2.c 和 EX2_Traffic_Emergency.c，主要工程文件为 Exp2_int1.uvproj、Exp2_int2.uvproj 和 Exp2_traffic.uvproj。Proteus 中由于器件库限制，采用与 8051 内核兼容的 AT89C51RD2 建立交通灯虚拟电路，用于替代真机中的 SST89E554RC 完成仿真验证。")
    add_body_paragraph(content_cell, "基础题 1 中，P1.0 和 P1.1 分别作为两路方波输出端；基础题 2 中，P1 用于驱动 8 个 LED，P3.2 和 P3.3 分别作为 INT0 和 INT1 外部中断输入；提高题中，P1.0~P1.5 分别连接南北红、南北黄、南北绿、东西红、东西黄、东西绿，P3.2 作为急救车请求中断输入。")

    add_heading(content_cell, "三、实验线路示图、程序算法流程图")
    add_body_paragraph(content_cell, "1. 实验线路示意图")
    circuit_img = EX2_IMG / "Circuit_Simulation_Diagram.png"
    if circuit_img.exists():
        add_picture(content_cell, circuit_img, 13.8, "图1  实验二 Proteus 交通灯仿真电路图")
    else:
        add_center_note(content_cell, "图1  实验二 Proteus 交通灯仿真电路图（待补）")
    add_body_paragraph(content_cell, "2. 程序算法说明")
    add_body_paragraph(content_cell, "（1）基础题 1：初始化 TH0、TL0、TH1、TL1 为 0xF800，设置 TMOD = 0x11 使定时器 0、1 以 16 位方式工作，设置 TCON = 0x50 启动两个定时器，设置 IE = 0x8A 打开总中断和两个定时器中断。在 timer0 和 timer1 中断服务程序中分别对 P1.0 和 P1.1 取反并重装定时器，从而形成两路稳定方波。")
    add_body_paragraph(content_cell, "（2）基础题 2：主程序把 P1 清零后配置 INT0、INT1 为下降沿触发，并打开 EX0、EX1 与 EA。INT0 服务程序中让 P1 在 0xFF 和 0x00 之间交替切换 4 次，实现全灯闪烁；INT1 服务程序中令变量 i 从 0x03 开始每次循环左移 2 位，对应 P1 输出 0x03、0x0C、0x30、0xC0 的序列，实现两灯一组流水显示。")
    add_body_paragraph(content_cell, "（3）提高题：采用状态机思想，把交通灯划分为 NS_GREEN、NS_YELLOW、EW_GREEN、EW_YELLOW 四种状态，并通过 show_state() 把状态映射为 P1 的具体输出值。Timer0 中断产生软件秒节拍，主循环依据 remain_seconds 完成状态保持与切换。INT0 到来时不直接长时间处理，只置位 emergency_req；主循环检测到该请求后保存当前 state 与 remain_seconds，令系统进入 10 s 全红保护，结束后恢复到中断前状态与剩余时间。")
    add_body_paragraph(content_cell, "（4）实验二对应的各题 Mermaid 流程图已整理为单独文件《实验二算法流程图_mermaid.md》，包括基础题 1、基础题 2、提高题交通灯主状态机流程图，以及提高题 Proteus 接线示意图，便于后续直接渲染插入实验报告。")

    add_heading(content_cell, "四、实验调试步骤、实验数据记录及实验结果")
    add_body_paragraph(content_cell, "1. 实验调试步骤")
    add_body_paragraph(content_cell, "（1）分别建立 EX2_Int1、EX2_INT2 和 Exp2_traffic 工程，添加源文件并编译生成目标文件和 HEX 文件。")
    add_body_paragraph(content_cell, "（2）在基础题 1 中，通过 Logic Analyzer 观察 Wave1 与 Wave2 的波形，验证两个定时器中断是否都能稳定工作。")
    add_body_paragraph(content_cell, "（3）在基础题 2 中，通过 Port3 窗口对 P3.2 和 P3.3 制造 1→0→1 的下降沿，分别触发 INT0 和 INT1，并在 Port1 与 Logic Analyzer 中观察 LED 显示结果。")
    add_body_paragraph(content_cell, "（4）在提高题中，先在 Keil 中通过 Watch 观察 state、remain_seconds、emergency_req、emergency_active 等变量的变化，再在 Proteus 中搭建交通灯电路，加载 EX2_Traffic_Emergency.hex 进行联动验证。")
    add_body_paragraph(content_cell, "（5）调试过程中重点检查定时器重装值、时钟频率、INT0 按键极性、RST 网络极性以及保存恢复变量是否正确，从而确保程序现象与设计逻辑一致。")
    add_body_paragraph(content_cell, "2. 实验数据记录及实验结果")
    add_body_paragraph(content_cell, "（1）基础题 1：EX2_Int1.c 中 P1.0 和 P1.1 均输出规则方波。Logic Analyzer 中 Wave1 与 Wave2 呈现稳定翻转波形，说明 timer0_isr() 和 timer1_isr() 都能在各自定时器溢出后正确执行，定时器 0 和定时器 1 初始化无误。")
    ex21 = EX2_IMG / "Ex2_1_1.png"
    if ex21.exists():
        add_picture(content_cell, ex21, 13.5, "图2  基础题1中 P1.0 与 P1.1 的方波输出结果")
    else:
        add_center_note(content_cell, "图2  基础题1中 P1.0 与 P1.1 的方波输出结果（待补）")
    add_body_paragraph(content_cell, "（2）基础题 2：EX2_INT2.c 中，触发 INT0 后 P1 在 0xFF 和 0x00 之间交替变化，对应 8 个 LED 全亮全灭闪烁 4 次；触发 INT1 后，P1 依次呈现 0x03、0x0C、0x30、0xC0 的循环序列，对应两灯一组由右向左流水显示。仿真图中已经能观察到这两类现象。")
    ex22 = EX2_IMG / "Ex2_1_2.png"
    if ex22.exists():
        add_picture(content_cell, ex22, 13.5, "图3  基础题2中外部中断触发后的 LED 输出结果")
    else:
        add_center_note(content_cell, "图3  基础题2中外部中断触发后的 LED 输出结果（待补）")
    add_body_paragraph(content_cell, "（3）提高题：EX2_Traffic_Emergency.c 中，正常状态下交通灯按‘南北绿 8 s—南北黄 2 s—东西绿 8 s—东西黄 2 s’循环运行，对应 P1 分别输出 0x0C、0x0A、0x21、0x11；当 INT0 急救车请求到来后，emergency_req 被置位，主循环保存当前 state 与 remain_seconds，并使系统立即切换为全红状态，即 P1 = 0x09，保持 10 s。全红结束后，程序恢复到中断前状态与剩余时间，而不是回到固定初始态。")
    ex2mid = EX2_IMG / "Ex2_mid.png"
    if ex2mid.exists():
        add_picture(content_cell, ex2mid, 13.8, "图4  提高题在 Keil 中的状态变量与端口变化结果")
    else:
        add_center_note(content_cell, "图4  提高题在 Keil 中的状态变量与端口变化结果（待补）")
    add_body_paragraph(content_cell, "（4）Proteus 仿真结果：采用 AT89C51RD2 搭建虚拟电路后，加载 EX2_Traffic_Emergency.hex，可观察到南北和东西方向的红黄绿灯正常轮转；按下 INT0 对应按键后，系统立即全红，保持设定时间后恢复到请求前状态。该结果与 Keil 中的变量变化和端口输出一致，说明提高题设计能够同时满足逻辑正确性和电路可实现性。")
    add_body_paragraph(content_cell, "（5）调试中还观察到，仿真时钟频率和定时器阈值会显著影响软件秒的体感长短。通过结合 TH0/TL0 装值、Xtal 频率以及 tick_50ms 比较阈值的计算，可以解释为什么同一套程序在不同仿真频率下会出现快慢差异。这一步有助于建立对‘时钟—定时器—时间尺度’关系的更准确认识。")

    add_heading(content_cell, "五、实验总结与心得")
    add_body_paragraph(content_cell, "通过本次实验，我对 51 单片机中断系统的理解明显更加具体。基础题让我比较扎实地掌握了定时器中断、外部中断以及 Port 口观察方法，不再只是会写寄存器配置，而是能够把定时器重装值、中断触发方式和端口现象联系起来分析。")
    add_body_paragraph(content_cell, "提高题进一步说明，单片机程序设计不能只停留在单个中断服务程序或单个端口输出上，而需要从系统层面思考：什么时候由中断置位请求、什么时候由主循环处理复杂逻辑、如何保存现场、怎样在异常事件结束后恢复原状态。emergency_req、emergency_active、saved_state 和 saved_remain 这组变量的设计，就是把“快速响应”和“系统稳定恢复”统一起来的关键。")
    add_body_paragraph(content_cell, "从调试体验来看，Keil 和 Proteus 的结合非常有价值。Keil 更适合检查变量、寄存器和中断响应过程，Proteus 更适合验证电路连接与整体现象。通过两者结合，实验完成了从代码逻辑到电路现象的闭环验证，也让我对后续更复杂嵌入式控制系统的设计过程有了更清晰的认识。")

    doc.save(str(OUTPUT))


if __name__ == '__main__':
    build_report()
